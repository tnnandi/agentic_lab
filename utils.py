import os
import re
from datetime import datetime
from docx import Document
from duckduckgo_search import DDGS
from bs4 import BeautifulSoup
import requests
from llm_utils import query_llm
import prompts


def save_output(report, code, execution_result, iteration):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = f"./output_agent/output_{timestamp}"
    os.makedirs(output_dir, exist_ok=True)

    # save report
    if report:
        doc = Document()
        doc.add_heading(f"Research Report - Iteration {iteration + 1}", level=1)
        doc.add_paragraph(report)
        report_file = os.path.join(output_dir, f"research_report_iteration_{iteration + 1}.docx")
        doc.save(report_file)

    # save code
    if code:
        code_file = os.path.join(output_dir, f"code_iteration_{iteration + 1}.py")
        with open(code_file, "w") as f:
            f.write(code)

    # save execution results
    if execution_result: 
        result_file = os.path.join(output_dir, f"execution_result_{iteration + 1}.txt")
        with open(result_file, "w") as f:
            f.write(execution_result)

    print(f"\nOutputs saved for iteration {iteration + 1} in {output_dir}")

# function to clean up the report to conform to professional standards
def clean_report(text):
    """
    Removes LLM reasoning, markdown, and formatting artifacts from report output.
    """
    # Remove <think>...</think> blocks
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)

    # Remove markdown-style headers and separators
    text = re.sub(r"^\s*#+\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*-{3,}\s*$", "", text, flags=re.MULTILINE)

    # Remove extra leading/trailing whitespace
    return text.strip()
    
def extract_code_only(text):
    """
    Extracts clean Python code from an LLM response by:
    - Removing <think>...</think> blocks
    - Removing markdown code fences (```python ... ```)
    - Returning only executable code with inline comments
    """
    # Remove <think>...</think> sections
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)

    # Extract the code inside ```python ... ``` blocks, if they exist
    match = re.search(r"```(?:python)?\n(.*?)```", text, re.DOTALL)
    code = match.group(1).strip() if match else text.strip()

    return code
    
def quick_duckduckgo_search(query, max_results=3):
    print(f"Performing quick DuckDuckGo search for: '{query}'")
    try:
        with DDGS() as ddgs:
            results = ddgs.text(query)
            top_results = list(results)[:max_results]

        raw_text = "\n\n".join(
            f"{i+1}. {r['title']}\n    {r['href']}\n    {r.get('body', '').strip()}"
            for i, r in enumerate(top_results)
        )

        print("Summarizing top search results...\n")

        summary_prompt = prompts.get_quick_search_summary_prompt(query, raw_text)

        raw_summary = query_llm(summary_prompt).strip()
        # Remove <think>...</think> block if present
        summary = re.sub(r"<think>.*?</think>", "", raw_summary, flags=re.DOTALL).strip()

        return f"Answer:\n{summary}\n\nBased on DuckDuckGo Search Results:\n\n{raw_text}"

    except Exception as e:
        return f"DuckDuckGo search failed: {e}"